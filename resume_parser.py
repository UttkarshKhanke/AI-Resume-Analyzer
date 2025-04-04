import os
from PyPDF2 import PdfReader
import docx

# PDF Extraction
def extract_text_from_pdf(filepath):
    text = ""
    try:
        reader = PdfReader(filepath)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text.strip() + " \n"  # Add space to avoid merged words
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text.strip()

# DOCX Extraction
def extract_text_from_docx(filepath):
    text = ""
    try:
        doc = docx.Document(filepath)
        print(f"Reading DOCX: {filepath}")

        if not doc.paragraphs:
            print("No paragraphs found in DOCX.")

        for para in doc.paragraphs:
            print("Paragraph:", para.text)  # Debug
            text += para.text.strip() + "\n"

    except Exception as e:
        print(f"Error reading DOCX: {e}")
    
    print("Final extracted DOCX text:\n", text)  # Debug full content
    return text.strip()


# Extract skills from text
def extract_skills(text):
    skill_keywords = [
        "python", "java", "c++", "sql", "html", "css", "javascript", "react",
        "node.js", "machine learning", "deep learning", "data analysis",
        "tensorflow", "pandas", "numpy", "matplotlib", "git", "docker",
        "kubernetes", "linux", "flask", "django", "communication", "teamwork",
        "problem solving", "leadership"
    ]

    text_lower = text.lower()
    found_skills = []

    for skill in skill_keywords:
        if skill.lower() in text_lower:
            # Capitalize properly for display
            if skill == "c++":
                found_skills.append("C++")
            elif skill == "sql":
                found_skills.append("SQL")
            elif skill == "html":
                found_skills.append("HTML")
            elif skill == "css":
                found_skills.append("CSS")
            elif skill == "javascript":
                found_skills.append("JavaScript")
            elif skill == "node.js":
                found_skills.append("Node.js")
            elif skill == "tensorflow":
                found_skills.append("TensorFlow")
            elif skill == "pandas":
                found_skills.append("Pandas")
            elif skill == "numpy":
                found_skills.append("NumPy")
            elif skill == "matplotlib":
                found_skills.append("Matplotlib")
            elif skill == "flask":
                found_skills.append("Flask")
            elif skill == "django":
                found_skills.append("Django")
            elif skill == "git":
                found_skills.append("Git")
            elif skill == "docker":
                found_skills.append("Docker")
            elif skill == "kubernetes":
                found_skills.append("Kubernetes")
            elif skill == "linux":
                found_skills.append("Linux")
            elif skill == "python":
                found_skills.append("Python")
            elif skill == "java":
                found_skills.append("Java")
            else:
                # Title-case the rest
                found_skills.append(skill.title())

    print("Extracted Skills:", found_skills)
    return found_skills


def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        print("Unsupported file format:", ext)
        return ""
